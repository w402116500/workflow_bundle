from __future__ import annotations

import argparse
import csv
import json
import os
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor
from PIL import Image

from core.code_image_renderer import (
    build_bundled_font_candidates,
    prepare_code_image_layout,
    render_prepared_code_lines_image,
    split_code_image_lines,
)
from core.document_format import (
    NUMBER_TOKEN_RE,
    normalize_caption_text,
    normalize_inline_reference_text,
    normalize_internal_figure_number,
    resolve_document_format,
)

WORKSPACE_DEFAULT_ROOT = Path(__file__).resolve().parents[2]
DEFAULT_CONFIG_PATH = None


LINE_SPACING_PT = 23
TITLE_SPACING_HALF_LINE = LINE_SPACING_PT / 2
BODY_FIRST_LINE_INDENT_PT = 24


FIG_RE = re.compile(r"!\[(?P<alt>.*?)\]\((?P<path>.*?)\)")
HEADING_RE = re.compile(r"^(?P<level>#{1,6})\s+(?P<text>.+?)\s*$")
TABLE_SEP_RE = re.compile(r"^\|\s*:?-{1,}:?\s*(\|\s*:?-{1,}:?\s*)+\|\s*$")
TABLE_ROW_RE = re.compile(r"^\|.*\|\s*$")
TABLE_CAPTION_RE = re.compile(rf"^表{NUMBER_TOKEN_RE}\s+.+")
FIG_CAPTION_RE = re.compile(rf"^图{NUMBER_TOKEN_RE}\s+.+")
FIG_PLACEHOLDER_RE = re.compile(r"^（配图占位，终稿插入图(?P<figs>.+?)）\s*$")
FIG_HIDDEN_MARKER_RE = re.compile(r"^<!--\s*figure:\s*(?P<figs>.+?)\s*-->\s*$")
REF_ENTRY_RE = re.compile(r"^\[(\d+)\]\s*")
CITE_RE = re.compile(r"\[(\d+)\]")
TRAILING_URL_RE = re.compile(r"\s+(?:https?|ftp)://\S+\s*$", re.IGNORECASE)


DEFAULT_CHAPTER_ORDER = [
    "01-绪论.md",
    "02-系统开发工具及技术介绍.md",
    "03-需求分析.md",
    "04-系统设计.md",
    "05-系统实现.md",
    "06-系统测试.md",
    "07-结论与展望.md",
    "08-致谢.md",
    "REFERENCES.md",
]
DEFAULT_KEYWORDS_CN = "FISCO BCOS；区块链；健康档案；访问控制；存证审计"
DEFAULT_KEYWORDS_EN = "FISCO BCOS; blockchain; health record; access control; notarization and audit"
CODE_RENDER_FONT_ENV_VAR = "THESIS_CODE_SCREENSHOT_FONT"
CODE_RENDER_FONT_SIZE_PX = 13
CODE_RENDER_BORDER_PX = 1
CODE_RENDER_IMAGE_PAD_X_PX = 12
CODE_RENDER_IMAGE_PAD_Y_PX = 10
CODE_RENDER_LINE_PAD_PX = 0
CODE_RENDER_PARAGRAPH_LEFT_INDENT_PT = 10
CODE_RENDER_MM_PER_PX = 0.25
CODE_RENDER_MAX_DISPLAY_WIDTH_MM = 148.0
CODE_RENDER_MAX_DISPLAY_HEIGHT_MM = 135.0
CODE_RENDER_MAX_CONTENT_WIDTH_PX = max(
    240,
    int(round(CODE_RENDER_MAX_DISPLAY_WIDTH_MM / CODE_RENDER_MM_PER_PX))
    - (CODE_RENDER_IMAGE_PAD_X_PX * 2)
    - (CODE_RENDER_BORDER_PX * 2),
)
CODE_RENDER_FIXED_CANVAS_WIDTH_PX = CODE_RENDER_MAX_CONTENT_WIDTH_PX
CODE_RENDER_BUNDLED_FONT_RELATIVE_PATHS = [
    Path("assets/fonts/sarasa-mono-sc/SarasaMonoSC-Regular.ttf"),
    Path("assets/fonts/siyuan-heiti/SourceHanSansSC-Regular-2.otf"),
]
CODE_RENDER_FONT_PATH_CANDIDATES = [
    Path("C:/Windows/Fonts/consola.ttf"),
    Path("C:/Windows/Fonts/consolab.ttf"),
    Path("C:/Windows/Fonts/Consolas.ttf"),
    Path("/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc"),
    Path("/usr/share/fonts/truetype/noto/NotoSansMono-Regular.ttf"),
    Path("/usr/share/fonts/truetype/noto/NotoMono-Regular.ttf"),
    Path("C:/Windows/Fonts/msyh.ttc"),
    Path("C:/Windows/Fonts/msyhbd.ttc"),
    Path("C:/Windows/Fonts/Deng.ttf"),
    Path("C:/Windows/Fonts/simsun.ttc"),
    Path("C:/Windows/Fonts/simsunb.ttf"),
    Path("C:/Windows/Fonts/Arial.ttf"),
]
CODE_BLOCK_TEXT_STYLE_PRESETS: dict[str, dict[str, object]] = {
    "plain-paper": {
        "font_cn": "SimSun",
        "font_en": "Times New Roman",
        "size_pt": 10.5,
        "line_spacing": "single",
        "left_indent_pt": 10.0,
        "space_before_pt": 6.0,
        "space_after_pt": 6.0,
        "tab_size": 4,
    },
    "mono-block": {
        "font_cn": "SimSun",
        "font_en": "Consolas",
        "size_pt": 10.5,
        "line_spacing": "single",
        "left_indent_pt": 10.0,
        "space_before_pt": 6.0,
        "space_after_pt": 6.0,
        "tab_size": 4,
    },
}


def _resolve_default_config_path(config_path: Path | None = None) -> Path:
    if config_path:
        return Path(config_path).resolve()

    if WORKSPACE_DEFAULT_ROOT.name == "workflow_bundle":
        active_pointer = WORKSPACE_DEFAULT_ROOT / "workflow" / "configs" / "active_workspace.json"
    else:
        active_pointer = WORKSPACE_DEFAULT_ROOT / "workflow_bundle" / "workflow" / "configs" / "active_workspace.json"

    if not active_pointer.exists():
        raise FileNotFoundError(
            "No active workspace configured. Run `python3 workflow_bundle/tools/cli.py set-active-workspace --config <workspace.json>` first."
        )

    payload = json.loads(active_pointer.read_text(encoding="utf-8"))
    raw_path = str(payload.get("config_path") or "").strip()
    if not raw_path:
        raise FileNotFoundError(f"Active workspace pointer is invalid: {active_pointer}")
    resolved = Path(raw_path)
    if not resolved.is_absolute():
        base_root = WORKSPACE_DEFAULT_ROOT.parent if WORKSPACE_DEFAULT_ROOT.name == "workflow_bundle" else WORKSPACE_DEFAULT_ROOT
        resolved = (base_root / resolved).resolve()
    if not resolved.exists():
        raise FileNotFoundError(f"Active workspace config does not exist: {resolved}")
    return resolved


@dataclass
class FigureItem:
    caption: str
    source_path: Path
    processed_path: Path
    show_caption: bool = True
    inserted_page: int | None = None


@dataclass
class BuildSettings:
    workspace_root: Path
    config_path: Path | None
    input_dir: Path
    diagram_dir: Path
    output_dir: Path
    output_docx: Path
    processed_img_dir: Path
    figure_log: Path
    reference_file: str
    abstract_cn_file: str
    abstract_en_file: str
    chapter_order: list[str]
    default_keywords_cn: str
    default_keywords_en: str
    figure_map: dict[str, tuple[str, Path]]
    document_format: dict[str, object]


def _resolve_path(base: Path, raw_path: str) -> Path:
    path = Path(raw_path)
    if path.is_absolute():
        return path
    return (base / path).resolve()


def _resolve_optional_path(base: Path, raw_path: str | None, fallback: Path) -> Path:
    if not raw_path:
        return fallback
    return _resolve_path(base, raw_path)


def _resolve_output_member(base: Path, raw_path: str | None, fallback_name: str) -> Path:
    if not raw_path:
        return base / fallback_name
    path = Path(raw_path)
    if path.is_absolute():
        return path
    return base / path


def _default_settings(workspace_root: Path | None = None) -> BuildSettings:
    root = (workspace_root or WORKSPACE_DEFAULT_ROOT).resolve()
    diagram_dir = root / "docs" / "images"
    output_dir = root / "word_output"
    return BuildSettings(
        workspace_root=root,
        config_path=None,
        input_dir=root / "polished_v3",
        diagram_dir=diagram_dir,
        output_dir=output_dir,
        output_docx=output_dir / "论文_润色版_规范排版_含配图_v29.docx",
        processed_img_dir=output_dir / "processed_images",
        figure_log=output_dir / "figure_insert_log_v29.csv",
        reference_file="REFERENCES.md",
        abstract_cn_file="00-摘要.md",
        abstract_en_file="00-Abstract.md",
        chapter_order=list(DEFAULT_CHAPTER_ORDER),
        default_keywords_cn=DEFAULT_KEYWORDS_CN,
        default_keywords_en=DEFAULT_KEYWORDS_EN,
        figure_map={
            "3.1": ("图3.1 系统用例图", diagram_dir / "image.png"),
            "4.1": ("图4.1 系统总体架构图", diagram_dir / "image-5.png"),
            "4.2": ("图4.2 数据库E-R图", diagram_dir / "image-4.png"),
            "4.3": ("图4.3 上传确认链上存证流程图", diagram_dir / "image-2.png"),
            "4.4": ("图4.4 授权撤销权限校验流程图", diagram_dir / "image-6.png"),
            "4.5": ("图4.5 带权限查询与审计追溯流程图", diagram_dir / "image-3.png"),
            "5.1": ("图5.1 系统功能结构图", diagram_dir / "image-1.png"),
        },
        document_format=resolve_document_format({}),
    )


def _load_settings(config_path: Path | None) -> BuildSettings:
    if config_path is None or not config_path.exists():
        return _default_settings()

    raw = json.loads(config_path.read_text(encoding="utf-8"))
    workspace_root_raw = raw.get("workspace_root")
    workspace_root = (
        _resolve_path(config_path.parent, workspace_root_raw)
        if workspace_root_raw
        else WORKSPACE_DEFAULT_ROOT.resolve()
    )
    defaults = _default_settings(workspace_root)
    build_cfg = raw.get("build", {})
    defaults_cfg = raw.get("defaults", {})
    document_format = resolve_document_format(raw)

    output_dir = _resolve_optional_path(workspace_root, build_cfg.get("output_dir"), defaults.output_dir)
    raw_figure_map = raw.get("figure_map")
    settings = BuildSettings(
        workspace_root=workspace_root,
        config_path=config_path.resolve(),
        input_dir=_resolve_optional_path(workspace_root, build_cfg.get("input_dir"), defaults.input_dir),
        diagram_dir=_resolve_optional_path(workspace_root, build_cfg.get("diagram_dir"), defaults.diagram_dir),
        output_dir=output_dir,
        output_docx=_resolve_output_member(output_dir, build_cfg.get("output_docx"), defaults.output_docx.name),
        processed_img_dir=_resolve_output_member(
            output_dir,
            build_cfg.get("processed_image_dir"),
            defaults.processed_img_dir.name,
        ),
        figure_log=_resolve_output_member(output_dir, build_cfg.get("figure_log"), defaults.figure_log.name),
        reference_file=build_cfg.get("reference_file", defaults.reference_file),
        abstract_cn_file=build_cfg.get("abstract_cn_file", defaults.abstract_cn_file),
        abstract_en_file=build_cfg.get("abstract_en_file", defaults.abstract_en_file),
        chapter_order=list(build_cfg.get("chapter_order") or defaults.chapter_order),
        default_keywords_cn=defaults_cfg.get("keywords_cn", defaults.default_keywords_cn),
        default_keywords_en=defaults_cfg.get("keywords_en", defaults.default_keywords_en),
        # Do not inherit legacy template figures when a workspace config is present.
        # Configured workspaces should opt in only to the figure assets they actually stage.
        figure_map={},
        document_format=document_format,
    )

    for fig_num, fig_cfg in (raw_figure_map or {}).items():
        if not isinstance(fig_cfg, dict):
            continue
        caption = fig_cfg.get("caption")
        path = fig_cfg.get("path")
        if not caption or not path:
            continue
        settings.figure_map[fig_num] = (caption, _resolve_path(workspace_root, path))

    return settings


def _extract_figure_number(text: str) -> str | None:
    match = re.match(rf"^图(?P<num>{NUMBER_TOKEN_RE})\s+.+$", text.strip())
    if not match:
        return None
    return normalize_internal_figure_number(match.group("num"))


def _caption_mentions_figure(text: str | None, fig_num: str) -> bool:
    if not text or not fig_num:
        return False
    extracted = _extract_figure_number(text.strip())
    return extracted == normalize_internal_figure_number(fig_num)


def _next_nonempty_line(lines: list[str], start_index: int) -> str:
    for idx in range(start_index, len(lines)):
        candidate = lines[idx].strip()
        if candidate:
            return candidate
    return ""


def _prev_nonempty_line(lines: list[str], start_index: int) -> str:
    for idx in range(start_index, -1, -1):
        candidate = lines[idx].strip()
        if candidate:
            return candidate
    return ""


SETTINGS = _default_settings()


def _activate_settings(settings: BuildSettings) -> None:
    global SETTINGS
    settings.output_dir.mkdir(parents=True, exist_ok=True)
    settings.processed_img_dir.mkdir(parents=True, exist_ok=True)
    for stale in settings.processed_img_dir.glob("codeblock_*"):
        if stale.is_file():
            stale.unlink()
    SETTINGS = settings


def _document_format() -> dict[str, object]:
    return SETTINGS.document_format


def _body_format() -> dict[str, object]:
    return dict(_document_format().get("body", {}))


def _code_block_render_mode() -> str:
    raw = str(_document_format().get("code_blocks", {}).get("render_mode", "image") or "image").strip().lower()
    return "text" if raw == "text" else "image"


def _code_block_format() -> dict[str, object]:
    raw = dict(_document_format().get("code_blocks", {}))
    text_style = str(raw.get("text_style", "plain-paper") or "plain-paper").strip().lower()
    preset = dict(CODE_BLOCK_TEXT_STYLE_PRESETS.get(text_style, CODE_BLOCK_TEXT_STYLE_PRESETS["plain-paper"]))
    for key, value in raw.items():
        if key in {"render_mode", "text_style"}:
            continue
        preset[key] = value
    preset["text_style"] = text_style
    return preset


def _line_spacing_pt() -> float:
    body = _body_format()
    return float(body.get("line_spacing_pt", LINE_SPACING_PT))


def _title_spacing_pt() -> float:
    return _line_spacing_pt() / 2


def _body_first_line_indent_pt() -> float:
    body = _body_format()
    return float(body.get("first_line_indent_pt", BODY_FIRST_LINE_INDENT_PT))


def _alignment_from_name(name: str | None) -> WD_ALIGN_PARAGRAPH:
    mapping = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return mapping.get(str(name or "left").lower(), WD_ALIGN_PARAGRAPH.LEFT)


def _set_style_rfonts(style, east_asia: str, ascii_font: str):
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for k in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        try:
            rfonts.attrib.pop(qn(k), None)
        except Exception:
            pass
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:cs"), ascii_font)


def _set_run_rfonts(run, east_asia: str, ascii_font: str):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for k in ("w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"):
        try:
            rfonts.attrib.pop(qn(k), None)
        except Exception:
            pass
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:cs"), ascii_font)


def _force_rpr_color_black(rpr) -> None:
    color = rpr.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        rpr.append(color)
    for attr in ("w:themeColor", "w:themeTint", "w:themeShade"):
        try:
            color.attrib.pop(qn(attr), None)
        except Exception:
            pass
    color.set(qn("w:val"), "000000")


def _set_style_color_black(style) -> None:
    style.font.color.rgb = RGBColor(0, 0, 0)
    _force_rpr_color_black(style.element.get_or_add_rPr())


def _set_run_color_black(run) -> None:
    run.font.color.rgb = RGBColor(0, 0, 0)
    _force_rpr_color_black(run._element.get_or_add_rPr())


def _force_style_element_color_black_by_id(doc: Document, style_id: str) -> None:
    for style_el in doc.styles.element.findall(qn("w:style")):
        if style_el.get(qn("w:styleId")) != style_id:
            continue
        rpr = style_el.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            style_el.append(rpr)
        _force_rpr_color_black(rpr)
        return


def _clear_paragraph(paragraph) -> None:
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def _add_field_run(paragraph, instruction: str) -> None:
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run_begin = paragraph.add_run()
    run_begin._element.append(begin)

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    run_instr = paragraph.add_run()
    run_instr._element.append(instr_text)
    _set_run_rfonts(run_instr, east_asia="SimSun", ascii_font="Times New Roman")
    _set_run_color_black(run_instr)

    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    run_sep = paragraph.add_run()
    run_sep._element.append(separate)

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run_end = paragraph.add_run()
    run_end._element.append(end)


def _apply_header_footer(doc: Document) -> None:
    header_footer = dict(_document_format().get("header_footer", {}))
    if not header_footer.get("enabled"):
        return

    header_text = str(header_footer.get("header_text", "") or "").strip()
    header_font_cn = str(header_footer.get("header_font_cn", "SimSun"))
    header_font_en = str(header_footer.get("header_font_en", "Times New Roman"))
    header_size = float(header_footer.get("header_size_pt", 10.5))
    footer_size = float(header_footer.get("footer_size_pt", 10.5))
    footer_line_spacing = _line_spacing_pt()

    for section in doc.sections:
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False

        header_para = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        _clear_paragraph(header_para)
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_para.paragraph_format.space_before = Pt(0)
        header_para.paragraph_format.space_after = Pt(0)
        header_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        header_para.paragraph_format.line_spacing = Pt(footer_line_spacing)
        if header_text:
            run = header_para.add_run(header_text)
            run.font.size = Pt(header_size)
            _set_run_rfonts(run, east_asia=header_font_cn, ascii_font=header_font_en)
            _set_run_color_black(run)

        footer_para = section.footer.paragraphs[0] if section.footer.paragraphs else section.footer.add_paragraph()
        _clear_paragraph(footer_para)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.paragraph_format.space_before = Pt(0)
        footer_para.paragraph_format.space_after = Pt(0)
        footer_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        footer_para.paragraph_format.line_spacing = Pt(footer_line_spacing)

        text_run = footer_para.add_run("第")
        text_run.font.size = Pt(footer_size)
        _set_run_rfonts(text_run, east_asia="SimSun", ascii_font="Times New Roman")
        _set_run_color_black(text_run)
        _add_field_run(footer_para, " PAGE ")

        text_run = footer_para.add_run("页 共")
        text_run.font.size = Pt(footer_size)
        _set_run_rfonts(text_run, east_asia="SimSun", ascii_font="Times New Roman")
        _set_run_color_black(text_run)
        _add_field_run(footer_para, " NUMPAGES ")

        text_run = footer_para.add_run("页")
        text_run.font.size = Pt(footer_size)
        _set_run_rfonts(text_run, east_asia="SimSun", ascii_font="Times New Roman")
        _set_run_color_black(text_run)


def _configure_page(doc: Document):
    page = dict(_document_format().get("page", {}))
    for section in doc.sections:
        section.page_width = Mm(float(page.get("width_mm", 210.0)))
        section.page_height = Mm(float(page.get("height_mm", 297.0)))
        section.orientation = WD_ORIENT.PORTRAIT
        section.top_margin = Mm(float(page.get("top_margin_mm", 25.0)))
        section.bottom_margin = Mm(float(page.get("bottom_margin_mm", 20.0)))
        section.left_margin = Mm(float(page.get("left_margin_mm", 25.0)))
        section.right_margin = Mm(float(page.get("right_margin_mm", 20.0)))


def _configure_styles(doc: Document):
    body = _body_format()
    headings = dict(_document_format().get("headings", {}))
    toc = dict(_document_format().get("toc", {}))
    captions = dict(_document_format().get("captions", {}))
    line_spacing_pt = _line_spacing_pt()

    normal = doc.styles["Normal"]
    normal.font.size = Pt(float(body.get("size_pt", 12.0)))
    normal.font.bold = False
    _set_style_rfonts(normal, east_asia=str(body.get("font_cn", "SimSun")), ascii_font=str(body.get("font_en", "Times New Roman")))
    _set_style_color_black(normal)
    ppf = normal.paragraph_format
    ppf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    ppf.line_spacing = Pt(line_spacing_pt)
    ppf.space_before = Pt(0)
    ppf.space_after = Pt(0)

    def add_heading_style(name: str, level_key: str, size_pt: float, align: WD_ALIGN_PARAGRAPH):
        cfg = dict(headings.get(level_key, {}))
        if name in doc.styles:
            st = doc.styles[name]
        else:
            st = doc.styles.add_style(name, 1)
        st.font.size = Pt(float(cfg.get("size_pt", size_pt)))
        st.font.bold = bool(cfg.get("bold", False))
        east_asia = str(cfg.get("font_cn", "SimHei"))
        ascii_font = str(cfg.get("font_en", "Times New Roman"))
        _set_style_rfonts(st, east_asia=east_asia, ascii_font=ascii_font)
        _set_style_color_black(st)
        pf = st.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line_spacing_pt)
        pf.space_before = Pt(float(cfg.get("space_before_pt", _title_spacing_pt())))
        pf.space_after = Pt(float(cfg.get("space_after_pt", _title_spacing_pt())))
        pf.alignment = _alignment_from_name(str(cfg.get("align", ""))) if cfg.get("align") else align
        try:
            linked = doc.styles[f"{name} Char"]
        except Exception:
            linked = None
        if linked is not None:
            linked.font.bold = bool(cfg.get("bold", False))
            _set_style_rfonts(linked, east_asia=east_asia, ascii_font=ascii_font)
            _set_style_color_black(linked)
        return st

    add_heading_style("Heading 1", "1", 18, WD_ALIGN_PARAGRAPH.LEFT)
    add_heading_style("Heading 2", "2", 15, WD_ALIGN_PARAGRAPH.LEFT)
    add_heading_style("Heading 3", "3", 14, WD_ALIGN_PARAGRAPH.LEFT)
    add_heading_style("Heading 4", "4", 12, WD_ALIGN_PARAGRAPH.LEFT)
    for style_id in ("Heading1", "Heading2", "Heading3", "Heading4", "Heading1Char", "Heading2Char", "Heading3Char", "Heading4Char"):
        _force_style_element_color_black_by_id(doc, style_id)

    def _configure_toc_style(name: str, level_key: str, east_asia: str, size_pt: float):
        cfg = dict(toc.get(level_key, {}))
        try:
            st = doc.styles[name]
        except Exception:
            return
        st.font.size = Pt(float(cfg.get("size_pt", size_pt)))
        st.font.bold = bool(cfg.get("bold", False))
        _set_style_rfonts(st, east_asia=str(cfg.get("font_cn", east_asia)), ascii_font=str(cfg.get("font_en", "Times New Roman")))
        _set_style_color_black(st)
        pf = st.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line_spacing_pt)
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)

    _configure_toc_style("TOC 1", "1", east_asia="SimHei", size_pt=14)
    _configure_toc_style("TOC 2", "2", east_asia="SimSun", size_pt=12)
    _configure_toc_style("TOC 3", "3", east_asia="SimSun", size_pt=12)

    cap = doc.styles.add_style("FigureCaption", 1) if "FigureCaption" not in doc.styles else doc.styles["FigureCaption"]
    fig_cap = dict(captions.get("figure", {}))
    cap.font.size = Pt(float(fig_cap.get("size_pt", 12.0)))
    cap.font.bold = bool(fig_cap.get("bold", True))
    _set_style_rfonts(cap, east_asia=str(fig_cap.get("font_cn", "SimHei")), ascii_font=str(fig_cap.get("font_en", "Times New Roman")))
    _set_style_color_black(cap)
    cap_pf = cap.paragraph_format
    cap_pf.alignment = _alignment_from_name(str(fig_cap.get("align", "center")))
    cap_pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    cap_pf.line_spacing = Pt(float(fig_cap.get("line_spacing_pt", line_spacing_pt)))
    cap_pf.space_before = Pt(0)
    cap_pf.space_after = Pt(0)

    tcap = doc.styles.add_style("TableCaption", 1) if "TableCaption" not in doc.styles else doc.styles["TableCaption"]
    tbl_cap = dict(captions.get("table", {}))
    tcap.font.size = Pt(float(tbl_cap.get("size_pt", 12.0)))
    tcap.font.bold = bool(tbl_cap.get("bold", True))
    _set_style_rfonts(tcap, east_asia=str(tbl_cap.get("font_cn", "SimHei")), ascii_font=str(tbl_cap.get("font_en", "Times New Roman")))
    _set_style_color_black(tcap)
    tcap_pf = tcap.paragraph_format
    tcap_pf.alignment = _alignment_from_name(str(tbl_cap.get("align", "center")))
    tcap_pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    tcap_pf.line_spacing = Pt(float(tbl_cap.get("line_spacing_pt", line_spacing_pt)))
    tcap_pf.space_before = Pt(0)
    tcap_pf.space_after = Pt(0)


def _add_toc(doc: Document):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(LINE_SPACING_PT)
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(0)

    container = p._p

    r_begin = OxmlElement("w:r")
    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")
    fc_begin.set(qn("w:dirty"), "true")
    r_begin.append(fc_begin)
    container.append(r_begin)

    r_instr = OxmlElement("w:r")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r_instr.append(instr)
    container.append(r_instr)

    r_sep = OxmlElement("w:r")
    fc_sep = OxmlElement("w:fldChar")
    fc_sep.set(qn("w:fldCharType"), "separate")
    r_sep.append(fc_sep)
    container.append(r_sep)

    hint_run = p.add_run("右键单击此处更新目录")
    hint_run.font.size = Pt(12)
    hint_run.font.bold = False
    _set_run_rfonts(hint_run, east_asia="SimSun", ascii_font="Times New Roman")
    _set_run_color_black(hint_run)

    r_end = OxmlElement("w:r")
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    r_end.append(fc_end)
    container.append(r_end)


def _apply_body_paragraph_format(paragraph, indent: bool = True):
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(_line_spacing_pt())
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Pt(_body_first_line_indent_pt()) if indent else Pt(0)


def _fix_zoom_percent(doc: Document):
    settings = doc.settings.element
    zoom = settings.find(qn("w:zoom"))
    if zoom is None:
        zoom = OxmlElement("w:zoom")
        zoom.set(qn("w:val"), "bestFit")
        settings.insert(0, zoom)
    if zoom.get(qn("w:percent")) is None:
        zoom.set(qn("w:percent"), "100")


def _ensure_update_fields_on_open(doc: Document):
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        update.set(qn("w:val"), "true")
        settings.append(update)
    else:
        update.set(qn("w:val"), "true")


def _add_page_break(doc: Document):
    doc.add_page_break()


def _add_section_break(doc: Document):
    doc.add_section(WD_SECTION_START.NEW_PAGE)


def _render_code_images(code_lines: list[str], output_path: Path, font_size: int = CODE_RENDER_FONT_SIZE_PX) -> list[Path]:
    font_candidates: list[Path] = []
    custom_font = os.environ.get(CODE_RENDER_FONT_ENV_VAR, "").strip()
    if custom_font:
        font_candidates.append(Path(custom_font))
    font_candidates.extend(build_bundled_font_candidates(Path(__file__), CODE_RENDER_BUNDLED_FONT_RELATIVE_PATHS))
    font_candidates.extend(CODE_RENDER_FONT_PATH_CANDIDATES)

    layout = prepare_code_image_layout(
        code_lines,
        font_candidates=font_candidates,
        font_size=font_size,
        max_content_width_px=CODE_RENDER_MAX_CONTENT_WIDTH_PX,
    )

    max_image_height_px = max(240, int(round(CODE_RENDER_MAX_DISPLAY_HEIGHT_MM / CODE_RENDER_MM_PER_PX)))
    chunks = split_code_image_lines(
        layout,
        padding_y=CODE_RENDER_IMAGE_PAD_Y_PX,
        line_pad=CODE_RENDER_LINE_PAD_PX,
        border_px=CODE_RENDER_BORDER_PX,
        max_image_height_px=max_image_height_px,
    )

    rendered_paths: list[Path] = []
    stem = output_path.stem
    suffix = output_path.suffix or ".png"
    for idx, chunk_lines in enumerate(chunks, start=1):
        chunk_path = output_path if len(chunks) == 1 else output_path.with_name(f"{stem}_part{idx}{suffix}")
        render_prepared_code_lines_image(
            chunk_lines,
            chunk_path,
            font=layout.font,
            padding_x=CODE_RENDER_IMAGE_PAD_X_PX,
            padding_y=CODE_RENDER_IMAGE_PAD_Y_PX,
            line_pad=CODE_RENDER_LINE_PAD_PX,
            border_px=CODE_RENDER_BORDER_PX,
            fixed_canvas_width_px=CODE_RENDER_FIXED_CANVAS_WIDTH_PX,
        )
        rendered_paths.append(chunk_path)
    return rendered_paths


def _normalize_code_line_text(text: str, tab_size: int) -> str:
    expanded = str(text).replace("\t", " " * max(1, tab_size))
    leading_spaces = len(expanded) - len(expanded.lstrip(" "))
    if leading_spaces > 0:
        expanded = ("\u00A0" * leading_spaces) + expanded[leading_spaces:]
    return expanded or "\u00A0"


def _add_text_code_block(doc: Document, code_lines: list[str]) -> None:
    code_cfg = _code_block_format()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.left_indent = Pt(float(code_cfg.get("left_indent_pt", 10.0)))
    pf.first_line_indent = Pt(0)
    pf.space_before = Pt(float(code_cfg.get("space_before_pt", 6.0)))
    pf.space_after = Pt(float(code_cfg.get("space_after_pt", 6.0)))
    if str(code_cfg.get("line_spacing", "single") or "single").strip().lower() == "single":
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    else:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(float(code_cfg.get("line_spacing_pt", _line_spacing_pt())))

    font_cn = str(code_cfg.get("font_cn", "SimSun") or "SimSun")
    font_en = str(code_cfg.get("font_en", "Times New Roman") or "Times New Roman")
    font_size = Pt(float(code_cfg.get("size_pt", 10.5)))
    tab_size = int(code_cfg.get("tab_size", 4) or 4)

    for index, line in enumerate(code_lines):
        run = p.add_run(_normalize_code_line_text(line, tab_size))
        run.font.size = font_size
        run.font.bold = False
        run.font.italic = False
        run.font.name = font_en
        _set_run_rfonts(run, east_asia=font_cn, ascii_font=font_en)
        _set_run_color_black(run)
        if index < len(code_lines) - 1:
            run.add_break()


def _add_title(doc: Document, text: str, font: str, bold: bool, style: str | None = None):
    heading_cfg = dict(_document_format().get("headings", {}).get("1", {}))
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(_line_spacing_pt())
    p.paragraph_format.space_before = Pt(float(heading_cfg.get("space_before_pt", _title_spacing_pt())))
    p.paragraph_format.space_after = Pt(float(heading_cfg.get("space_after_pt", _title_spacing_pt())))
    r = p.add_run(text)
    r.font.size = Pt(float(heading_cfg.get("size_pt", 18.0)))
    r.font.bold = bool(heading_cfg.get("bold", bold))
    r.font.name = str(heading_cfg.get("font_en", "Times New Roman"))
    _set_run_rfonts(r, east_asia=str(heading_cfg.get("font_cn", font)), ascii_font=str(heading_cfg.get("font_en", "Times New Roman")))
    _set_run_color_black(r)
    return p


def _add_keywords_cn(doc: Document, keywords: str):
    body = _body_format()
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(_line_spacing_pt())
    r1 = p.add_run("关键词：")
    r1.font.size = Pt(float(body.get("size_pt", 12.0)))
    r1.font.bold = False
    _set_run_rfonts(r1, east_asia="SimHei", ascii_font=str(body.get("font_en", "Times New Roman")))
    _set_run_color_black(r1)
    r2 = p.add_run(normalize_inline_reference_text(keywords, _document_format()))
    r2.font.size = Pt(float(body.get("size_pt", 12.0)))
    r2.font.bold = False
    _set_run_rfonts(r2, east_asia=str(body.get("font_cn", "SimSun")), ascii_font=str(body.get("font_en", "Times New Roman")))
    _set_run_color_black(r2)


def _add_keywords_en(doc: Document, keywords: str):
    body = _body_format()
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(_line_spacing_pt())
    r1 = p.add_run("Key words: ")
    r1.font.size = Pt(float(body.get("size_pt", 12.0)))
    r1.font.bold = True
    r1.font.name = str(body.get("font_en", "Times New Roman"))
    _set_run_color_black(r1)
    r2 = p.add_run(keywords)
    r2.font.size = Pt(float(body.get("size_pt", 12.0)))
    r2.font.bold = False
    r2.font.name = str(body.get("font_en", "Times New Roman"))
    _set_run_color_black(r2)


def _sanitize_heading(text: str) -> str:
    return text.replace("（初稿）", "").replace("(初稿)", "").strip()


def _is_center_title(text: str) -> bool:
    normalized = re.sub(r"\s+", "", text)
    return normalized in {"结论", "致谢", "参考文献", "附录"}

def _format_cite_no(num: str) -> str:
    return f"[{num}]"


def _add_reference_paragraph(doc: Document, text: str, bookmark_state: dict):
    m = REF_ENTRY_RE.match(text.strip())
    ref_no = m.group(1) if m else None
    p = doc.add_paragraph()
    _apply_body_paragraph_format(p, indent=False)
    if ref_no:
        rest = _strip_reference_url(text.strip()[m.end():].lstrip())
        r_no = p.add_run(_format_cite_no(ref_no))
        r_no.font.size = Pt(10.5)
        r_no.font.bold = False
        _set_run_rfonts(r_no, east_asia="SimSun", ascii_font="Times New Roman")
        _add_bookmark_on_run(r_no, f"ref_{ref_no}", bookmark_state)

        if rest:
            r_rest = p.add_run(" " + rest)
            r_rest.font.size = Pt(10.5)
            r_rest.font.bold = False
            _set_run_rfonts(r_rest, east_asia="SimSun", ascii_font="Times New Roman")
    else:
        r = p.add_run(_strip_reference_url(text.strip()))
        r.font.size = Pt(10.5)
        r.font.bold = False
        _set_run_rfonts(r, east_asia="SimSun", ascii_font="Times New Roman")


def _strip_reference_url(text: str) -> str:
    cleaned = TRAILING_URL_RE.sub("", text).rstrip()
    return cleaned

def _add_bookmark(paragraph, name: str, bookmark_state: dict | None = None):
    state = bookmark_state if bookmark_state is not None else getattr(_add_bookmark, "_state", None)
    if state is None:
        state = {"next_id": 1}
        setattr(_add_bookmark, "_state", state)
    bid = state["next_id"]
    state["next_id"] = bid + 1

    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))

    p = paragraph._p
    ppr = p.find(qn("w:pPr"))
    insert_at = 1 if ppr is not None else 0
    p.insert(insert_at, start)
    p.append(end)

def _add_bookmark_on_run(run, name: str, bookmark_state: dict):
    bid = bookmark_state["next_id"]
    bookmark_state["next_id"] = bid + 1

    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bid))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bid))

    r = run._r
    p = r.getparent()
    idx = p.index(r)
    p.insert(idx, start)
    p.insert(idx + 2, end)


def _make_superscript_rpr(font_size_half_pt: str = "21"):
    rpr = OxmlElement("w:rPr")

    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), "SimSun")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")
    rpr.append(rfonts)

    v = OxmlElement("w:vertAlign")
    v.set(qn("w:val"), "superscript")
    rpr.append(v)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "none")
    rpr.append(u)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    rpr.append(color)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), font_size_half_pt)
    rpr.append(sz)
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), font_size_half_pt)
    rpr.append(sz_cs)

    return rpr


def _add_ref_field(paragraph, bookmark_name: str, display_text: str):
    p = paragraph._p

    r1 = OxmlElement("w:r")
    r1.append(_make_superscript_rpr())
    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")
    r1.append(fc_begin)
    p.append(r1)

    r2 = OxmlElement("w:r")
    r2.append(_make_superscript_rpr())
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" REF {bookmark_name} \\h "
    r2.append(instr)
    p.append(r2)

    r3 = OxmlElement("w:r")
    r3.append(_make_superscript_rpr())
    fc_sep = OxmlElement("w:fldChar")
    fc_sep.set(qn("w:fldCharType"), "separate")
    r3.append(fc_sep)
    p.append(r3)

    r4 = OxmlElement("w:r")
    r4.append(_make_superscript_rpr())
    t = OxmlElement("w:t")
    t.text = display_text
    r4.append(t)
    p.append(r4)

    r5 = OxmlElement("w:r")
    r5.append(_make_superscript_rpr())
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    r5.append(fc_end)
    p.append(r5)


def _append_field(
    paragraph,
    instruction: str,
    *,
    display_text: str = "",
    font_size_pt: float = 12,
    east_asia: str = "SimSun",
    ascii_font: str = "Times New Roman",
    bold: bool = False,
) -> None:
    def _formatted_run(text: str | None = None):
        run = paragraph.add_run(text or "")
        run.font.size = Pt(font_size_pt)
        run.font.bold = bold
        _set_run_rfonts(run, east_asia=east_asia, ascii_font=ascii_font)
        _set_run_color_black(run)
        return run

    begin_run = _formatted_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    begin_run._r.append(fld_begin)

    instr_run = _formatted_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "
    instr_run._r.append(instr)

    separate_run = _formatted_run()
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(fld_sep)

    if display_text:
        _formatted_run(display_text)

    end_run = _formatted_run()
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    end_run._r.append(fld_end)


def _make_field_rpr(
    *,
    font_size_pt: float = 12,
    east_asia: str = "SimSun",
    ascii_font: str = "Times New Roman",
    bold: bool = False,
):
    rpr = OxmlElement("w:rPr")

    bold_el = OxmlElement("w:b")
    bold_el.set(qn("w:val"), "1" if bold else "0")
    rpr.append(bold_el)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    rpr.append(color)

    size_half_pt = str(int(round(font_size_pt * 2)))
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), size_half_pt)
    rpr.append(sz)

    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), size_half_pt)
    rpr.append(sz_cs)

    rfonts = OxmlElement("w:rFonts")
    rfonts.set(qn("w:eastAsia"), east_asia)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:cs"), ascii_font)
    rpr.append(rfonts)

    return rpr


def _append_simple_field(
    paragraph,
    instruction: str,
    *,
    display_text: str = "",
    font_size_pt: float = 12,
    east_asia: str = "SimSun",
    ascii_font: str = "Times New Roman",
    bold: bool = False,
) -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instruction)
    fld.set(qn("w:dirty"), "true")

    run = OxmlElement("w:r")
    run.append(
        _make_field_rpr(
            font_size_pt=font_size_pt,
            east_asia=east_asia,
            ascii_font=ascii_font,
            bold=bold,
        )
    )
    if display_text:
        text = OxmlElement("w:t")
        text.text = display_text
        run.append(text)
    fld.append(run)
    paragraph._p.append(fld)


def _add_superscript_text(paragraph, text: str):
    run = paragraph.add_run(text)
    run.font.superscript = True
    run.font.size = Pt(10.5)
    run.font.size = Pt(10.5)
    _set_run_rfonts(run, east_asia="SimSun", ascii_font="Times New Roman")
    _set_run_color_black(run)


def _clear_paragraph_content(paragraph) -> None:
    p = paragraph._p
    for child in list(p):
        if child.tag == qn("w:pPr"):
            continue
        p.remove(child)


def _configure_page_number_footers(doc: Document) -> None:
    for idx, section in enumerate(doc.sections):
        footer = section.footer
        if idx > 0 and footer.is_linked_to_previous:
            footer.is_linked_to_previous = False
            footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        _clear_paragraph_content(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.first_line_indent = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        paragraph.paragraph_format.line_spacing = 1.0
        _append_simple_field(
            paragraph,
            "PAGE \\* MERGEFORMAT",
            display_text="1",
            font_size_pt=10.5,
            east_asia="Times New Roman",
            ascii_font="Times New Roman",
        )


def _add_text_with_cite_links(paragraph, text: str, ref_nums: set[str]):
    # Inline markdown markers are not part of the required thesis format.
    # Strip them here as a safety net (Markdown sources are also cleaned).
    body = _body_format()
    body_cn_font = str(body.get("font_cn", "SimSun"))
    body_en_font = str(body.get("font_en", "Times New Roman"))
    s = normalize_inline_reference_text(text.strip().replace("`", "").replace("**", ""), _document_format())
    if not s:
        return
    pos = 0
    for m in CITE_RE.finditer(s):
        start, end = m.span()
        if start > pos:
            r = paragraph.add_run(s[pos:start])
            _set_run_rfonts(r, east_asia=body_cn_font, ascii_font=body_en_font)
            _set_run_color_black(r)
        num = m.group(1)
        cite_text = _format_cite_no(num)
        if num in ref_nums:
            _add_ref_field(paragraph, f"ref_{num}", cite_text)
        else:
            _add_superscript_text(paragraph, cite_text)
        pos = end
    if pos < len(s):
        r = paragraph.add_run(s[pos:])
        _set_run_rfonts(r, east_asia=body_cn_font, ascii_font=body_en_font)
        _set_run_color_black(r)


def _set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)

def _set_table_borders_half_pt(table, color: str = "000000"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    def _set_edge(name: str):
        el = borders.find(qn(f"w:{name}"))
        if el is None:
            el = OxmlElement(f"w:{name}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        _set_edge(edge)


def _set_border(el, val: str, size: str = "0", color: str = "000000"):
    el.set(qn("w:val"), val)
    if val == "nil":
        for attr in ("w:sz", "w:space", "w:color"):
            try:
                el.attrib.pop(qn(attr), None)
            except Exception:
                pass
        return
    el.set(qn("w:sz"), size)
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), color)


def _set_cell_border(cell, edge: str, val: str, size: str = "0", color: str = "000000"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    border = tc_borders.find(qn(f"w:{edge}"))
    if border is None:
        border = OxmlElement(f"w:{edge}")
        tc_borders.append(border)
    _set_border(border, val=val, size=size, color=color)


def _clear_table_style(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    style = tbl_pr.find(qn("w:tblStyle"))
    if style is not None:
        tbl_pr.remove(style)


def _apply_three_line_table(table, color: str = "000000"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    edge_defs = {
        "top": ("single", "8"),
        "bottom": ("single", "8"),
        "left": ("nil", "0"),
        "right": ("nil", "0"),
        "insideH": ("nil", "0"),
        "insideV": ("nil", "0"),
    }
    for edge, (val, size) in edge_defs.items():
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        _set_border(el, val=val, size=size, color=color)

    if table.rows:
        for cell in table.rows[0].cells:
            _set_cell_border(cell, "top", val="single", size="8", color=color)
            _set_cell_border(cell, "bottom", val="single", size="8", color=color)
            for edge in ("left", "right"):
                _set_cell_border(cell, edge, val="nil", color=color)

    for row in table.rows[1:]:
        for cell in row.cells:
            _set_cell_border(cell, "top", val="nil", color=color)
            _set_cell_border(cell, "bottom", val="nil", color=color)
            for edge in ("left", "right"):
                _set_cell_border(cell, edge, val="nil", color=color)

    if len(table.rows) > 1:
        for cell in table.rows[-1].cells:
            _set_cell_border(cell, "bottom", val="single", size="8", color=color)


def _process_image(src: Path) -> Path:
    if not src.exists():
        raise FileNotFoundError(str(src))
    dst = SETTINGS.processed_img_dir / src.name
    if src.resolve() == dst.resolve():
        return dst
    with Image.open(src) as im:
        im = im.convert("RGB")
        im.save(dst, format="PNG", optimize=True)
    return dst


def _resolve_markdown_image_source(md_path: Path, rel_path: str) -> Path:
    candidate = (md_path.parent / rel_path).resolve()
    if candidate.exists():
        return candidate

    basename = Path(rel_path).name
    search_roots = [
        SETTINGS.diagram_dir,
        SETTINGS.workspace_root / "images",
        SETTINGS.workspace_root / "docs",
        SETTINGS.output_dir / "processed_images",
        SETTINGS.processed_img_dir,
    ]

    for root in search_roots:
        direct = root / basename
        if direct.exists():
            return direct

    for root in search_roots:
        if not root.exists():
            continue
        matches = list(root.rglob(basename))
        if matches:
            return matches[0]

    raise FileNotFoundError(str(candidate))


def _normalize_markdown_image_line(line: str) -> str:
    normalized = (
        line.replace(r"\[", "[")
        .replace(r"\]", "]")
        .replace(r"\(", "(")
        .replace(r"\)", ")")
    )
    normalized = re.sub(r"\s+null\)$", ")", normalized)
    return normalized


def _is_code_screenshot_image(src: Path, alt: str) -> bool:
    path_parts = {part.lower() for part in src.parts}
    alt_text = alt.strip()
    return "code_screenshots" in path_parts or "代码截图" in alt_text


def _resolve_numbered_figure_override(markdown_src: Path, alt: str, pending_fig_caption: str | None) -> tuple[str | None, Path | None]:
    candidates = []
    if pending_fig_caption and pending_fig_caption.strip():
        candidates.append(pending_fig_caption.strip())
    if alt.strip():
        candidates.append(alt.strip())
    for candidate in candidates:
        if not FIG_CAPTION_RE.match(candidate):
            continue
        fig_num = _extract_figure_number(candidate)
        if not fig_num:
            continue
        mapped = SETTINGS.figure_map.get(fig_num)
        if not mapped:
            continue
        cap_text, mapped_src = mapped
        if _caption_mentions_figure(candidate, fig_num):
            cap_text = candidate
        return normalize_caption_text(cap_text, "figure", _document_format()), mapped_src
    return None, None


def _add_image_with_caption(doc: Document, fig: FigureItem):
    section = doc.sections[0]
    avail_width_mm = float((section.page_width - section.left_margin - section.right_margin) / Mm(1))
    avail_height_mm = float((section.page_height - section.top_margin - section.bottom_margin) / Mm(1))

    max_width_mm = min(150.0, avail_width_mm)
    max_height_mm = min(150.0, max(80.0, avail_height_mm - 40.0))

    with Image.open(fig.processed_path) as im:
        w_px, h_px = im.size
    aspect = w_px / max(1, h_px)
    width_mm = min(max_width_mm, max_height_mm * aspect)
    height_mm = width_mm / max(1e-6, aspect)
    if height_mm > max_height_mm:
        height_mm = max_height_mm
        width_mm = height_mm * aspect

    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p_img.paragraph_format.line_spacing = 1.0
    p_img.paragraph_format.keep_together = True
    p_img.paragraph_format.keep_with_next = fig.show_caption
    p_img.paragraph_format.space_before = Pt(12)
    p_img.paragraph_format.space_after = Pt(0 if fig.show_caption else 12)
    run = p_img.add_run()
    run.add_picture(str(fig.processed_path), width=Mm(width_mm), height=Mm(height_mm))

    if fig.show_caption and fig.caption.strip():
        p_cap = doc.add_paragraph(normalize_caption_text(fig.caption, "figure", _document_format()), style="FigureCaption")
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.keep_together = True
        p_cap.paragraph_format.space_before = Pt(0)
        p_cap.paragraph_format.space_after = Pt(0)


def _add_table(doc: Document, caption: str | None, headers: list[str], rows: list[list[str]]):
    if caption:
        doc.add_paragraph(normalize_caption_text((caption or "").replace("`", "").replace("**", ""), "table", _document_format()), style="TableCaption")

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _clear_table_style(table)

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = (h or "").replace("`", "").replace("**", "")

    for r in rows:
        row_cells = table.add_row().cells
        for i, v in enumerate(r):
            row_cells[i].text = (v or "").replace("`", "").replace("**", "")

    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                p.paragraph_format.line_spacing = Pt(_line_spacing_pt())
                for run in p.runs:
                    run.font.size = Pt(10.5)
                    run.font.bold = False
                    _set_run_rfonts(run, east_asia="SimSun", ascii_font="Times New Roman")
                    _set_run_color_black(run)
    _apply_three_line_table(table)


def _parse_md_and_add(doc: Document, md_path: Path, figures: list[FigureItem], page_state: dict, ref_nums: set[str], bookmark_state: dict):
    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    pending_table_caption: str | None = None
    pending_fig_caption: str | None = None
    code_block_index = 0
    inserted_reference_title = False

    if md_path.name == "REFERENCES.md":
        has_reference_heading = any(HEADING_RE.match((line or "").strip()) for line in lines)
        if not has_reference_heading:
            _add_title(doc, "参考文献", font="SimHei", bold=False, style="Heading 1")
            inserted_reference_title = True

    def add_paragraph(text: str):
        text = text.strip()
        if not text:
            return
        if md_path.name == "REFERENCES.md" and REF_ENTRY_RE.match(text):
            _add_reference_paragraph(doc, text, bookmark_state)
            return
        p = doc.add_paragraph()
        _apply_body_paragraph_format(p, indent=True)
        _add_text_with_cite_links(p, text, ref_nums)

    while i < len(lines):
        line = lines[i].rstrip()

        if md_path.name == "REFERENCES.md" and REF_ENTRY_RE.match(line.strip()):
            _add_reference_paragraph(doc, line.strip(), bookmark_state)
            i += 1
            continue

        if line.lstrip().startswith(("- ", "* ")):
            while i < len(lines) and lines[i].lstrip().startswith(("- ", "* ")):
                item = lines[i].lstrip()[2:].strip()
                if item:
                    p = doc.add_paragraph(style="List Bullet")
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    p.paragraph_format.line_spacing = Pt(_line_spacing_pt())
                    _add_text_with_cite_links(p, item, ref_nums)
                i += 1
            continue

        m_num = re.match(r"^\s*(\d+)\.\s+(.+)$", line)
        if m_num:
            while i < len(lines):
                m2 = re.match(r"^\s*(\d+)\.\s+(.+)$", lines[i].rstrip())
                if not m2:
                    break
                item = m2.group(2).strip()
                if item:
                    p = doc.add_paragraph(style="List Number")
                    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    p.paragraph_format.line_spacing = Pt(_line_spacing_pt())
                    _add_text_with_cite_links(p, item, ref_nums)
                i += 1
            continue

        if line.strip().startswith("```"):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            if i < len(lines) and lines[i].strip().startswith("```"):
                i += 1

            if code_lines:
                if _code_block_render_mode() == "text":
                    _add_text_code_block(doc, code_lines)
                else:
                    code_block_index += 1
                    img_name = f"codeblock_{md_path.stem}_{code_block_index}.png"
                    img_path = SETTINGS.processed_img_dir / img_name
                    rendered_paths = _render_code_images(code_lines, img_path)
                    for part_index, rendered_path in enumerate(rendered_paths, start=1):
                        p = doc.add_paragraph()
                        _apply_body_paragraph_format(p, indent=False)
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                        p.paragraph_format.line_spacing = 1.0
                        p.paragraph_format.left_indent = Pt(CODE_RENDER_PARAGRAPH_LEFT_INDENT_PT)
                        p.paragraph_format.space_before = Pt(6 if part_index == 1 else 2)
                        p.paragraph_format.space_after = Pt(2 if part_index < len(rendered_paths) else 6)
                        p.paragraph_format.keep_together = True
                        run = p.add_run()
                        with Image.open(rendered_path) as code_image:
                            width_px, height_px = code_image.size
                        width_mm = width_px * CODE_RENDER_MM_PER_PX
                        height_mm = height_px * CODE_RENDER_MM_PER_PX
                        scale = min(
                            1.0,
                            CODE_RENDER_MAX_DISPLAY_WIDTH_MM / max(width_mm, 1e-6),
                            CODE_RENDER_MAX_DISPLAY_HEIGHT_MM / max(height_mm, 1e-6),
                        )
                        width_mm *= scale
                        height_mm *= scale
                        run.add_picture(str(rendered_path), width=Mm(width_mm), height=Mm(height_mm))
            continue

        m_h = HEADING_RE.match(line)
        if m_h:
            level = len(m_h.group("level"))
            text = _sanitize_heading(m_h.group("text")).replace("`", "").replace("**", "")
            if md_path.name == "REFERENCES.md" and level >= 2:
                i += 1
                continue
            if md_path.name == "REFERENCES.md" and inserted_reference_title and level == 1 and text == "参考文献":
                i += 1
                continue
            if level == 1:
                heading_cfg = dict(_document_format().get("headings", {}).get("1", {}))
                p = doc.add_paragraph(style="Heading 1")
                r = p.add_run(text)
                _set_run_rfonts(r, east_asia=str(heading_cfg.get("font_cn", "SimHei")), ascii_font=str(heading_cfg.get("font_en", "Times New Roman")))
                _set_run_color_black(r)
                if _is_center_title(text):
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif level == 2:
                heading_cfg = dict(_document_format().get("headings", {}).get("2", {}))
                p = doc.add_paragraph(style="Heading 2")
                r = p.add_run(text)
                _set_run_rfonts(r, east_asia=str(heading_cfg.get("font_cn", "SimHei")), ascii_font=str(heading_cfg.get("font_en", "Times New Roman")))
                _set_run_color_black(r)
            elif level == 3:
                heading_cfg = dict(_document_format().get("headings", {}).get("3", {}))
                p = doc.add_paragraph(style="Heading 3")
                r = p.add_run(text)
                _set_run_rfonts(r, east_asia=str(heading_cfg.get("font_cn", "SimHei")), ascii_font=str(heading_cfg.get("font_en", "Times New Roman")))
                _set_run_color_black(r)
            elif level == 4:
                heading_cfg = dict(_document_format().get("headings", {}).get("4", {}))
                p = doc.add_paragraph(style="Heading 4")
                r = p.add_run(text)
                _set_run_rfonts(r, east_asia=str(heading_cfg.get("font_cn", "SimHei")), ascii_font=str(heading_cfg.get("font_en", "Times New Roman")))
                _set_run_color_black(r)
            else:
                add_paragraph(text)
            i += 1
            continue

        if TABLE_CAPTION_RE.match(line.strip()):
            pending_table_caption = normalize_caption_text(line.strip(), "table", _document_format())
            i += 1
            continue

        if FIG_CAPTION_RE.match(line.strip()):
            s = normalize_caption_text(line.strip(), "figure", _document_format())
            if "。" not in s and len(s) <= 40:
                fig_num = _extract_figure_number(s)
                next_nonempty = _next_nonempty_line(lines, i + 1)
                next_normalized = _normalize_markdown_image_line(next_nonempty) if next_nonempty else ""
                prev_nonempty = _prev_nonempty_line(lines, i - 1)
                prev_normalized = _normalize_markdown_image_line(prev_nonempty) if prev_nonempty else ""
                prev_fig_match = FIG_RE.search(prev_normalized) if prev_normalized else None
                prev_is_same_figure = False
                if prev_fig_match:
                    prev_alt = prev_fig_match.group("alt").strip()
                    prev_num = _extract_figure_number(prev_alt) if FIG_CAPTION_RE.match(prev_alt) else None
                    prev_is_same_figure = prev_num == fig_num
                if fig_num and prev_is_same_figure:
                    pending_fig_caption = None
                    i += 1
                    continue
                has_followup_figure = bool(
                    (next_nonempty and (FIG_HIDDEN_MARKER_RE.match(next_nonempty) or FIG_PLACEHOLDER_RE.match(next_nonempty)))
                    or (next_normalized and FIG_RE.search(next_normalized))
                )
                if fig_num and not has_followup_figure and fig_num in SETTINGS.figure_map:
                    cap_text, src = SETTINGS.figure_map[fig_num]
                    if _caption_mentions_figure(s, fig_num):
                        cap_text = s
                    processed = _process_image(src)
                    fig = FigureItem(caption=cap_text, source_path=src, processed_path=processed)
                    _add_image_with_caption(doc, fig)
                    figures.append(fig)
                    pending_fig_caption = None
                    i += 1
                    continue
                if fig_num and not has_followup_figure:
                    add_paragraph(s)
                    pending_fig_caption = None
                    i += 1
                    continue
                pending_fig_caption = s
                i += 1
                continue

        m_fig_marker = FIG_HIDDEN_MARKER_RE.match(line.strip()) or FIG_PLACEHOLDER_RE.match(line.strip())
        if m_fig_marker:
            figs_raw = m_fig_marker.group("figs")
            fig_nums = re.findall(r"\d+\.\d+", figs_raw)
            inserted_any = False
            missing = []
            for fig_num in fig_nums:
                if fig_num in SETTINGS.figure_map:
                    cap_text, src = SETTINGS.figure_map[fig_num]
                    if _caption_mentions_figure(pending_fig_caption, fig_num):
                        cap_text = pending_fig_caption
                    processed = _process_image(src)
                    fig = FigureItem(caption=cap_text, source_path=src, processed_path=processed)
                    _add_image_with_caption(doc, fig)
                    figures.append(fig)
                    inserted_any = True
                else:
                    missing.append(fig_num)
            pending_fig_caption = None
            if not inserted_any:
                add_paragraph(normalize_inline_reference_text(line.strip(), _document_format()))
            elif missing:
                remain = "、".join([f"图{normalize_internal_figure_number(n).replace('.', '-') if _document_format().get('numbering', {}).get('figure') == 'hyphen' else normalize_internal_figure_number(n)}" for n in missing])
                add_paragraph(f"（配图占位，终稿插入{remain}）")
            i += 1
            continue

        normalized_line = _normalize_markdown_image_line(line)
        m_fig = FIG_RE.search(normalized_line)
        if m_fig:
            alt = m_fig.group("alt").strip()
            rel = m_fig.group("path").strip()
            src = _resolve_markdown_image_source(md_path, rel)
            cap_override, src_override = _resolve_numbered_figure_override(src, alt, pending_fig_caption)
            if src_override is not None:
                src = src_override
            processed = _process_image(src)
            cap = cap_override or (pending_fig_caption.strip() if pending_fig_caption else alt)
            show_caption = not _is_code_screenshot_image(src, alt)
            fig = FigureItem(caption=cap, source_path=src, processed_path=processed, show_caption=show_caption)
            _add_image_with_caption(doc, fig)
            figures.append(fig)
            pending_fig_caption = None
            i += 1
            continue

        if TABLE_ROW_RE.match(line) and i + 1 < len(lines) and TABLE_SEP_RE.match(lines[i + 1].rstrip()):
            header_cells = [c.strip() for c in line.strip().strip("|").split("|")]
            i += 2
            body_rows: list[list[str]] = []
            while i < len(lines) and TABLE_ROW_RE.match(lines[i].rstrip()):
                row_cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                if len(row_cells) < len(header_cells):
                    row_cells += [""] * (len(header_cells) - len(row_cells))
                body_rows.append(row_cells[: len(header_cells)])
                i += 1
            _add_table(doc, pending_table_caption, header_cells, body_rows)
            pending_table_caption = None
            continue

        if not line.strip():
            i += 1
            continue

        para_lines = [line]
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip()
            if not nxt.strip():
                break
            if nxt.strip().startswith("```"):
                break
            normalized_nxt = _normalize_markdown_image_line(nxt)
            if HEADING_RE.match(nxt) or TABLE_CAPTION_RE.match(nxt.strip()) or FIG_RE.search(normalized_nxt):
                break
            if TABLE_ROW_RE.match(nxt) and i + 1 < len(lines) and TABLE_SEP_RE.match(lines[i + 1].rstrip()):
                break
            para_lines.append(nxt)
            i += 1

        add_paragraph(" ".join([p.strip() for p in para_lines]))


def _read_abstract_file(path: Path) -> tuple[str, str] | None:
    if not path.exists():
        return None
    text = path.read_text(encoding="utf-8").strip()
    if not text:
        return None
    lines = [l.strip() for l in text.splitlines() if l.strip()]
    keywords = ""
    body_lines: list[str] = []
    for l in lines:
        if HEADING_RE.match(l):
            continue
        if l.lower().startswith("key words:") or l.startswith("关键词："):
            if "：" in l:
                keywords = l.split("：", 1)[-1].strip()
            elif ":" in l:
                keywords = l.split(":", 1)[-1].strip()
            else:
                keywords = l.strip()
        else:
            body_lines.append(l)
    body = " ".join(body_lines).strip()
    return body, keywords


def build():
    doc = Document()
    _configure_page(doc)
    _configure_styles(doc)
    _fix_zoom_percent(doc)
    _ensure_update_fields_on_open(doc)

    page_state = {"page": 1}
    bookmark_state = {"next_id": 1}

    ref_nums: set[str] = set()
    ref_path = SETTINGS.input_dir / SETTINGS.reference_file
    if ref_path.exists():
        for ln in ref_path.read_text(encoding="utf-8").splitlines():
            m = REF_ENTRY_RE.match((ln or "").strip())
            if m:
                ref_nums.add(m.group(1))

    _add_title(doc, "摘  要", font="SimHei", bold=False, style="Heading 1")
    cn = _read_abstract_file(SETTINGS.input_dir / SETTINGS.abstract_cn_file)
    cn_body = cn[0] if cn else f"（请在此处填写中文摘要正文，行间距固定{int(_line_spacing_pt())}磅。）"
    cn_kw = cn[1] if cn and cn[1] else SETTINGS.default_keywords_cn
    p = doc.add_paragraph(normalize_inline_reference_text(cn_body, _document_format()))
    _apply_body_paragraph_format(p, indent=True)
    _add_keywords_cn(doc, cn_kw)
    _add_page_break(doc)
    page_state["page"] += 1

    _add_title(doc, "Abstract", font="Times New Roman", bold=True, style="Heading 1")
    en = _read_abstract_file(SETTINGS.input_dir / SETTINGS.abstract_en_file)
    en_body = en[0] if en else "(Fill in the English abstract here.)"
    en_kw = en[1] if en and en[1] else SETTINGS.default_keywords_en
    p = doc.add_paragraph(en_body)
    _apply_body_paragraph_format(p, indent=True)
    _add_keywords_en(doc, en_kw)
    _add_page_break(doc)
    page_state["page"] += 1

    _add_title(doc, "目  录", font="SimHei", bold=False)
    _add_toc(doc)
    _add_section_break(doc)
    page_state["page"] += 1

    md_files = [SETTINGS.input_dir / n for n in SETTINGS.chapter_order if (SETTINGS.input_dir / n).exists()]
    if not md_files:
        raise RuntimeError(f"no md files found in: {SETTINGS.input_dir}")

    figures: list[FigureItem] = []
    for idx, f in enumerate(md_files):
        _parse_md_and_add(doc, f, figures, page_state, ref_nums, bookmark_state)
        if idx != len(md_files) - 1:
            _add_page_break(doc)
            page_state["page"] += 1

    _configure_page_number_footers(doc)
    _apply_header_footer(doc)

    out_docx = SETTINGS.output_docx
    try:
        doc.save(str(out_docx))
    except PermissionError:
        out_docx = SETTINGS.output_dir / f"{SETTINGS.output_docx.stem}_v2{SETTINGS.output_docx.suffix}"
        doc.save(str(out_docx))

    with SETTINGS.figure_log.open("w", newline="", encoding="utf-8") as fp:
        w = csv.writer(fp)
        w.writerow(["figure_caption", "source_path", "processed_path", "inserted_page"])
        for f in figures:
            w.writerow([f.caption, str(f.source_path), str(f.processed_path), f.inserted_page or ""])
    return out_docx


def resolve_output_docx_path(config_path: Path | None = None, output_name: str | None = None) -> Path:
    settings = _load_settings(_resolve_default_config_path(config_path))
    if output_name:
        settings.output_docx = settings.output_dir / output_name
    return settings.output_docx


def _parse_args(argv: list[str] | None = None) -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Build thesis DOCX from a configured workspace.")
    parser.add_argument(
        "--config",
        type=Path,
        default=None,
        help="Path to workspace config JSON. Defaults to the active workspace pointer.",
    )
    parser.add_argument(
        "--output-name",
        help="Override the output DOCX filename inside the configured output directory.",
    )
    parser.add_argument(
        "--print-output-path",
        action="store_true",
        help="Print the resolved output DOCX path and exit without building.",
    )
    return parser.parse_args(argv)


def main(argv: list[str] | None = None) -> Path | None:
    args = _parse_args(argv)
    settings = _load_settings(_resolve_default_config_path(args.config))
    if args.output_name:
        settings.output_docx = settings.output_dir / args.output_name
    _activate_settings(settings)

    if args.print_output_path:
        print(SETTINGS.output_docx)
        return SETTINGS.output_docx

    return build()


if __name__ == "__main__":
    main()
